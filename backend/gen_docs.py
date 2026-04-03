"""
gen_docs.py — สร้าง 3 เอกสารหลังจาก optimize เสร็จ
  1. ใบปะหน้าข้อสอบ  (PDF via reportlab / docx)
  2. ปกซองข้อสอบ     (docx)
  3. ใบลงลายมือชื่อ   (docx)

ใช้: python gen_docs.py <schedule_id> <output_dir>
หรือ import ใช้ใน FastAPI router
"""

import io, os, re
from datetime import datetime
from typing import List, Optional

# ── python-docx ───────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Thai date helper ──────────────────────────────────────────
THAI_MONTHS = {
    1:"มกราคม",2:"กุมภาพันธ์",3:"มีนาคม",4:"เมษายน",
    5:"พฤษภาคม",6:"มิถุนายน",7:"กรกฎาคม",8:"สิงหาคม",
    9:"กันยายน",10:"ตุลาคม",11:"พฤศจิกายน",12:"ธันวาคม",
}
ENG_MONTHS = {
    1:"JAN",2:"FEB",3:"MAR",4:"APR",5:"MAY",6:"JUN",
    7:"JUL",8:"AUG",9:"SEP",10:"OCT",11:"NOV",12:"DEC",
}
THAI_DAYS = {
    0:"จันทร์",1:"อังคาร",2:"พุธ",3:"พฤหัสบดี",
    4:"ศุกร์",5:"เสาร์",6:"อาทิตย์",
}

def parse_exam_date(exam_date: str):
    """Parse '2026-03-23' → datetime"""
    try:
        return datetime.strptime(exam_date, "%Y-%m-%d")
    except:
        try:
            return datetime.strptime(exam_date, "%Y-%m-%d")
        except:
            return None

def thai_date_long(exam_date: str) -> str:
    """'2026-03-23' → 'วันจันทร์ที่ 23 มีนาคม 2569'"""
    d = parse_exam_date(exam_date)
    if not d:
        return exam_date
    thai_year = d.year + 543
    day_name = THAI_DAYS[d.weekday()]
    return f"วัน{day_name}ที่ {d.day} {THAI_MONTHS[d.month]} {thai_year}"

def thai_date_short(exam_date: str) -> str:
    """'2026-03-23' → 'จันทร์ 23 มี.ค. 2569'"""
    d = parse_exam_date(exam_date)
    if not d:
        return exam_date
    thai_year = d.year + 543
    short_months = {1:"ม.ค.",2:"ก.พ.",3:"มี.ค.",4:"เม.ย.",
                    5:"พ.ค.",6:"มิ.ย.",7:"ก.ค.",8:"ส.ค.",
                    9:"ก.ย.",10:"ต.ค.",11:"พ.ย.",12:"ธ.ค."}
    day_name = THAI_DAYS[d.weekday()]
    return f"{day_name} {d.day} {short_months[d.month]} {thai_year}"

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        if edge in kwargs:
            tag.set(qn('w:val'), kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:color'), kwargs[edge].get('color','000000'))
        else:
            tag.set(qn('w:val'), 'nil')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def bold_run(para, text, size_pt=None, underline=False, center=False):
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'TH Sarabun New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    if size_pt:
        run.font.size = Pt(size_pt)
    if underline:
        run.underline = True
    return run

def normal_run(para, text, size_pt=None, center=False):
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = 'TH Sarabun New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    if size_pt:
        run.font.size = Pt(size_pt)
    return run

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


# ════════════════════════════════════════════════════════════════
# 1. ใบปะหน้าข้อสอบ
# ════════════════════════════════════════════════════════════════

def generate_cover_page(
    course_id: str,
    course_name_th: str,
    section_no: str,
    exam_date: str,
    exam_time: str,
    room_name: str,
    exam_type: str = "final",  # midterm | final
    semester: str = "2",
    academic_year: str = "2568",
) -> bytes:
    """
    ใบปะหน้าข้อสอบ — 2 หน้า
    หน้า 1: คำชี้แจง
    หน้า 2: กรอบข้อมูล + ช่องกรอกชื่อ
    """
    doc = Document()

    # ตั้งค่าหน้า A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    exam_type_th = "ปลายภาค" if exam_type == "final" else "กลางภาค"
    date_long    = thai_date_long(exam_date)

    # ── หน้า 1: คำชี้แจง ──────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=60)
    bold_run(p, f"คณะรัฐศาสตร์และรัฐประศาสนศาสตร์  มหาวิทยาลัยเชียงใหม่", 16)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=120)
    bold_run(p, f"ข้อสอบ{exam_type_th} ภาคการศึกษาที่ {semester} ประจำปีการศึกษา {academic_year}", 16)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=60)
    bold_run(p, f"กระบวนวิชา  {course_id}   ตอน  {section_no}", 16)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=60)
    bold_run(p, f"สอบวันที่  {date_long}", 16)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=240)
    bold_run(p, f"เวลา  {exam_time}  น.   ห้องสอบ  {room_name}", 16)

    # เส้นคั่น
    p = doc.add_paragraph()
    set_para_spacing(p, before=60, after=120)
    pPr = p._element.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:color'), '000000')
    pb.append(bottom); pPr.append(pb)

    # หัวคำชี้แจง
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=120)
    bold_run(p, "คำชี้แจง", 16, underline=True)

    instructions = [
        "1.  ก่อนเข้าห้องสอบให้กรรมการคุมสอบประกาศให้นักศึกษาทราบว่า สิ่งของใดที่อนุญาตให้นำเข้าห้องสอบและสิ่งของใดที่ห้ามนำเข้าห้องสอบ หากตรวจพบว่านักศึกษาผู้ใดนำสิ่งของต้องห้ามเข้าห้องสอบ",
        "ถือเป็นการกระทำผิดฐานทุจริตในการสอบ",
        "2.  ให้นักศึกษาเขียนชื่อ – นามสกุล และรหัสประจำตัว ให้ชัดเจน",
        "3.  ก่อนเริ่มดำเนินการสอบให้กรรมการคุมสอบประกาศในห้องสอบให้นักศึกษาทราบเกี่ยวกับชื่อกระบวนวิชาที่จัดสอบ เวลาสอบ หมายเลขห้องสอบ และข้อปฏิบัติในการสอบ",
        "4.  ให้นักศึกษาปฏิบัติตามคำชี้แจงของกรรมการคุมสอบอย่างเคร่งครัด เมื่อต้องการสิ่งหนึ่งสิ่งใดในระหว่างการสอบให้ยกมือขึ้นและแจ้งความประสงค์ให้กรรมการคุมสอบทราบ",
        "5.  ให้นักศึกษาปฏิบัติตามข้อบังคับมหาวิทยาลัยเชียงใหม่ ว่าด้วยการสอบของนักศึกษา พ.ศ. 2564 อย่างเคร่งครัด",
        "6.  นักศึกษาที่ถูกกล่าวหาว่ากระทำการทุจริตในการสอบ ให้ดำเนินการทางวินัยนักศึกษาตามข้อบังคับมหาวิทยาลัยเชียงใหม่ ว่าด้วยวินัยและการดำเนินการทางวินัยนักศึกษา พ.ศ. 2564",
    ]

    for i, inst in enumerate(instructions):
        p = doc.add_paragraph()
        set_para_spacing(p, after=80)
        is_bold_line = i == 1  # "ถือเป็นการกระทำผิดฐานทุจริตในการสอบ"
        if is_bold_line:
            run = p.add_run(inst)
            run.bold = True; run.underline = True
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        else:
            run = p.add_run(inst)
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # ── หน้า 2: กรอบข้อมูล ──────────────────────────────────
    doc.add_page_break()

    # ช่องกรอกชื่อบนสุด
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=240)
    r = p.add_run("ชื่อ........................................นามสกุล.............................................รหัส..................................")
    r.font.name = 'TH Sarabun New'; r.font.size = Pt(14)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # กรอบข้อมูลวิชา
    tbl = doc.add_table(rows=5, cols=1)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)

    cell_data = [
        ("center", "คณะรัฐศาสตร์และรัฐประศาสนศาสตร์", 16, True),
        ("center", f"การสอบ{exam_type_th} ภาคการศึกษาที่ {semester} ปีการศึกษา {academic_year}", 15, False),
        ("center", f"กระบวนวิชา  {course_id}   ตอน  {section_no}", 15, True),
        ("center", f"สอบวันที่  {date_long}", 15, False),
        ("center", f"เวลา  {exam_time}  น.   ห้องสอบ  {room_name}", 15, False),
    ]

    for i, (align, text, size, is_bold) in enumerate(cell_data):
        cell = tbl.cell(i, 0)
        cell.width = Cm(16)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp = p.paragraph_format
        sp.space_before = Pt(8); sp.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(size)
        run.bold = is_bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # ช่องกรอกชื่อล่าง
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_para_spacing(p, before=120, after=0)
    r = p.add_run("ชื่อ....................................นามสกุล...........................................รหัส...............................")
    r.font.name = 'TH Sarabun New'; r.font.size = Pt(14)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════
# 2. ปกซองข้อสอบ
# ════════════════════════════════════════════════════════════════

def generate_envelope_cover(
    course_id: str,
    section_no: str,
    exam_date: str,
    exam_time: str,
    room_name: str,
    teacher_name: str,
    supervisors: List[dict],       # [{"name": ..., "slot_order": 1, "role_in_exam": "..."}]
    num_students: int,
    num_exam_sets: int,            # จำนวนชุดข้อสอบ
    exam_type: str = "final",
    semester: str = "2",
    academic_year: str = "2568",
    qr_image_path: Optional[str] = None,
    material_request: Optional[dict] = None,  # สิ่งที่ต้องการเพิ่มเติมจากอาจารย์
) -> bytes:
    """ปกซองข้อสอบ — 1 หน้าต่อ section"""
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21.0)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin = section.bottom_margin = Cm(1.5)

    exam_type_th = "ปลายภาค" if exam_type == "final" else "กลางภาค"
    exam_type_en = "Final Exams" if exam_type == "final" else "Midterm Exams"
    date_short   = thai_date_short(exam_date)

    # Header
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=40)
    bold_run(p, "คณะรัฐศาสตร์และรัฐประศาสนศาสตร์ มหาวิทยาลัยเชียงใหม่", 15)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=80)
    run = bold_run(p, f"การสอบ{exam_type_th}/{exam_type_en}", 15, underline=True)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, after=80)
    r = p.add_run(f"ภาคเรียนที่/Semester  {semester}   ปีการศึกษา/Academic Year  {academic_year}")
    r.font.name = 'TH Sarabun New'; r.font.size = Pt(14)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # ตารางข้อมูลวิชา
    info_tbl = doc.add_table(rows=3, cols=4)
    info_tbl.style = 'Table Grid'
    info_tbl.autofit = False
    widths = [Cm(4.5), Cm(3.5), Cm(3.5), Cm(4.5)]
    for i, w in enumerate(widths):
        for row in info_tbl.rows:
            row.cells[i].width = w

    def set_info_cell(row, col, label, value, bold_val=True):
        cell = info_tbl.cell(row, col)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{label}  ")
        r1.font.name = 'TH Sarabun New'; r1.font.size = Pt(13)
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        r2 = p.add_run(value)
        r2.font.name = 'TH Sarabun New'; r2.font.size = Pt(13); r2.bold = bold_val
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    set_info_cell(0, 0, "กระบวนวิชา/Course", course_id)
    set_info_cell(0, 2, "ตอน/Section", section_no)
    set_info_cell(1, 0, "วันสอบ/Date", date_short)
    set_info_cell(1, 2, "เวลา/Time", exam_time)

    # merge row2 cols 0-1 สำหรับห้อง
    cell_room = info_tbl.cell(2, 0).merge(info_tbl.cell(2, 1))
    p = cell_room.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("ห้องสอบ/Room  "); r1.font.name = 'TH Sarabun New'; r1.font.size = Pt(13)
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    r2 = p.add_run(room_name); r2.bold = True; r2.font.name = 'TH Sarabun New'; r2.font.size = Pt(13)
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # merge row2 cols 2-3 สำหรับอาจารย์
    cell_teacher = info_tbl.cell(2, 2).merge(info_tbl.cell(2, 3))
    p = cell_teacher.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("อาจารย์/Instructor  "); r1.font.name = 'TH Sarabun New'; r1.font.size = Pt(13)
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    r2 = p.add_run(teacher_name); r2.font.name = 'TH Sarabun New'; r2.font.size = Pt(13)
    r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # กรรมการคุมสอบ
    p = doc.add_paragraph(); set_para_spacing(p, before=120, after=60)
    bold_run(p, "กรรมการคุมสอบ/Invigilator", 14)

    sup_cols = 3 if qr_image_path else 2
    sup_tbl  = doc.add_table(rows=max(2, len(supervisors)), cols=sup_cols)
    sup_tbl.style = 'Table Grid'
    sup_tbl.autofit = False
    col_w = [Cm(8), Cm(6), Cm(2.5)] if qr_image_path else [Cm(9), Cm(7)]
    for i, w in enumerate(col_w[:sup_cols]):
        for row in sup_tbl.rows:
            if i < len(row.cells):
                row.cells[i].width = w

    for i, sup in enumerate(supervisors[:max(2, len(supervisors))]):
        if i >= len(sup_tbl.rows): break
        row = sup_tbl.rows[i]
        # ชื่อ
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{i+1}. {sup.get('name','')}")
        r.font.name = 'TH Sarabun New'; r.font.size = Pt(13)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        # ช่องลายเซ็น
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_before = Pt(4); p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run("......................................................")
        r2.font.name = 'TH Sarabun New'; r2.font.size = Pt(12)
        # QR code (row 0 เท่านั้น)
        if qr_image_path and sup_cols == 3 and i == 0:
            p3 = row.cells[2].paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = p3.add_run()
            run3.add_picture(qr_image_path, width=Cm(2.2))

    # จำนวนข้อสอบ
    p = doc.add_paragraph(); set_para_spacing(p, before=120, after=40)
    cnt_tbl = doc.add_table(rows=4, cols=3)
    cnt_tbl.style = 'Table Grid'
    cnt_tbl.autofit = False
    cw = [Cm(9), Cm(3), Cm(4)]
    for i, w in enumerate(cw):
        for row in cnt_tbl.rows:
            row.cells[i].width = w

    cnt_data = [
        (f"จำนวนข้อสอบ/Number of exam",  str(num_exam_sets), "ชุด"),
        ("Answer Sheet/กระดาษคำตอบ",       "",               "แผ่น"),
        ("Answer Paper/กระดาษเขียนตอบ",    "",               "แผ่น"),
        ("Note Paper/กระดาษทด",            "",               "แผ่น"),
    ]
    for i, (label, val, unit) in enumerate(cnt_data):
        row = cnt_tbl.rows[i]
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(3); p0.paragraph_format.space_after = Pt(3)
        r = p0.add_run(label); r.font.name = 'TH Sarabun New'; r.font.size = Pt(13)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        p1 = row.cells[1].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(val); r1.bold = True; r1.font.name = 'TH Sarabun New'; r1.font.size = Pt(13)
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        p2 = row.cells[2].paragraphs[0]
        p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(unit); r2.font.name = 'TH Sarabun New'; r2.font.size = Pt(13)
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # จำนวนนักศึกษา
    p = doc.add_paragraph(); set_para_spacing(p, before=80, after=40)
    stu_tbl = doc.add_table(rows=4, cols=3)
    stu_tbl.style = 'Table Grid'; stu_tbl.autofit = False
    for i, w in enumerate(cw):
        for row in stu_tbl.rows:
            row.cells[i].width = w

    stu_data = [
        ("จำนวนนักศึกษาที่ลงทะเบียน/Total of Students",     str(num_students), "คน"),
        ("จำนวนนักศึกษาที่เข้าสอบ/Number of attend students", "",               "คน"),
        ("จำนวนนักศึกษาที่ขาดสอบ/Number of absent students",  "",               "คน"),
        ("จำนวนนักศึกษาที่ถอน Q,W/Q,W",                      "",               "คน"),
    ]
    for i, (label, val, unit) in enumerate(stu_data):
        row = stu_tbl.rows[i]
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(3); p0.paragraph_format.space_after = Pt(3)
        r = p0.add_run(label); r.font.name = 'TH Sarabun New'; r.font.size = Pt(13)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        p1 = row.cells[1].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)
        r1 = p1.add_run(val); r1.bold = True; r1.font.name = 'TH Sarabun New'; r1.font.size = Pt(13)
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
        p2 = row.cells[2].paragraphs[0]
        p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
        r2 = p2.add_run(unit); r2.font.name = 'TH Sarabun New'; r2.font.size = Pt(13)
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # รายชื่อนักศึกษาขาดสอบ
    p = doc.add_paragraph(); set_para_spacing(p, before=80, after=40)
    bold_run(p, "รายชื่อนักศึกษาที่ขาดสอบ/Absent students list", 13)

    for i in range(1, 6):
        p = doc.add_paragraph(); set_para_spacing(p, after=60)
        r = p.add_run(f"{i}. " + "."*70)
        r.font.name = 'TH Sarabun New'; r.font.size = Pt(13)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════
# 3. ใบลงลายมือชื่อ
# ════════════════════════════════════════════════════════════════

def generate_attendance_sheet(
    course_id: str,
    section_no: str,
    exam_date: str,
    exam_time: str,
    room_name: str,
    teacher_name: str,
    students: List[dict],          # [{"student_id": ..., "student_name": ...}, ...]
    exam_type: str = "final",
    semester: str = "2",
    academic_year: str = "2568",
) -> bytes:
    """ใบลงลายมือชื่อผู้เข้าสอบ"""
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21.0)
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(1.5)

    exam_type_th = "ปลายภาค" if exam_type == "final" else "กลางภาค"
    date_long    = thai_date_long(exam_date)
    total        = len(students)

    def add_header(doc_obj, is_first=True):
        if not is_first:
            doc_obj.add_page_break()
        # ชื่อเอกสาร
        p = doc_obj.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=40)
        bold_run(p, "ใบลงลายมือชื่อผู้เข้าสอบ", 16)

        p = doc_obj.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=80)
        r = p.add_run("คณะรัฐศาสตร์และรัฐประศาสนศาสตร์ มหาวิทยาลัยเชียงใหม่")
        r.font.name = 'TH Sarabun New'; r.font.size = Pt(14)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

        r2_p = doc_obj.add_paragraph(); r2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(r2_p, after=80)
        r2 = r2_p.add_run(f"การสอบ{exam_type_th} ภาคการศึกษาที่ {semester} ปีการศึกษา {academic_year}")
        r2.font.name = 'TH Sarabun New'; r2.font.size = Pt(14)
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

        # ตารางข้อมูลวิชา
        info_tbl = doc_obj.add_table(rows=2, cols=4)
        info_tbl.style = 'Table Grid'; info_tbl.autofit = False
        iw = [Cm(4.5), Cm(4.0), Cm(3.5), Cm(5.0)]
        for ci, cw in enumerate(iw):
            for row in info_tbl.rows:
                row.cells[ci].width = cw

        def ic(r, c, lbl, val):
            cell = info_tbl.cell(r, c)
            p2 = cell.paragraphs[0]
            p2.paragraph_format.space_before = Pt(3); p2.paragraph_format.space_after = Pt(3)
            rx = p2.add_run(f"{lbl} "); rx.font.name = 'TH Sarabun New'; rx.font.size = Pt(13)
            rx._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
            ry = p2.add_run(val); ry.bold = True; ry.font.name = 'TH Sarabun New'; ry.font.size = Pt(13)
            ry._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

        ic(0, 0, "กระบวนวิชา", course_id)
        ic(0, 1, "ตอน", section_no)
        ic(0, 2, "ห้องสอบ", room_name)
        ic(0, 3, "", "")
        ic(1, 0, "วันสอบ", date_long)
        ic(1, 1, "", "")
        ic(1, 2, "เวลาสอบ", exam_time)
        ic(1, 3, "", "")

        # merge วันสอบ col 0-1
        info_tbl.cell(1, 0).merge(info_tbl.cell(1, 1))

        # จำนวน + อาจารย์
        p3 = doc_obj.add_paragraph(); set_para_spacing(p3, before=60, after=60)
        rr = p3.add_run(
            f"จำนวนผู้เข้าสอบ {total} คน (ทั้งหมด {total} คน)    "
            f"อาจารย์ประจำวิชา {teacher_name}"
        )
        rr.font.name = 'TH Sarabun New'; rr.font.size = Pt(13)
        rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    # header หน้าแรก
    add_header(doc, is_first=True)

    # ตารางรายชื่อ — 30 คนต่อหน้า
    ROWS_PER_PAGE = 30
    chunks = [students[i:i+ROWS_PER_PAGE] for i in range(0, max(1, len(students)), ROWS_PER_PAGE)]

    for page_idx, chunk in enumerate(chunks):
        if page_idx > 0:
            add_header(doc, is_first=False)

        sig_tbl = doc.add_table(rows=len(chunk)+1, cols=3)
        sig_tbl.style = 'Table Grid'; sig_tbl.autofit = False
        sw = [Cm(6.5), Cm(4.0), Cm(7.0)]
        for ci, cw in enumerate(sw):
            for row in sig_tbl.rows:
                row.cells[ci].width = cw

        # header row
        headers = ["ชื่อ - นามสกุล", "รหัสนักศึกษา", "ลายมือชื่อ"]
        hrow = sig_tbl.rows[0]
        for ci, h in enumerate(headers):
            p = hrow.cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(h); r.bold = True; r.font.name = 'TH Sarabun New'; r.font.size = Pt(13)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

        # รายชื่อ
        global_start = page_idx * ROWS_PER_PAGE
        for ri, stu in enumerate(chunk):
            row = sig_tbl.rows[ri + 1]
            row.height = Cm(0.85)

            # ชื่อ
            p0 = row.cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(2); p0.paragraph_format.space_after = Pt(2)
            r0 = p0.add_run(f"{global_start+ri+1}. {stu.get('student_name','')}")
            r0.font.name = 'TH Sarabun New'; r0.font.size = Pt(12)
            r0._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

            # รหัส
            p1 = row.cells[1].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_before = Pt(2); p1.paragraph_format.space_after = Pt(2)
            r1 = p1.add_run(str(stu.get('student_id','')))
            r1.font.name = 'TH Sarabun New'; r1.font.size = Pt(12)
            r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

            # ช่องลายเซ็น (ว่างเปล่า)
            row.cells[2].paragraphs[0].add_run("")

        # สรุปท้ายหน้า
        p_sum = doc.add_paragraph(); set_para_spacing(p_sum, before=60, after=0)
        start_n = global_start + 1
        end_n   = global_start + len(chunk)
        rs = p_sum.add_run(
            f"รวมหน้านี้  {len(chunk)}  คน   "
            f"(ลำดับที่ {start_n} – {end_n})   "
            f"รวมทั้งหมด  {total}  คน"
        )
        rs.font.name = 'TH Sarabun New'; rs.font.size = Pt(13); rs.bold = True
        rs._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════
# CLI test
# ════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    import os

    # ข้อมูลตัวอย่างจาก DB (126101 sec1)
    SAMPLE = {
        "course_id":      "126101",
        "course_name_th": "Introduction to International Relations",
        "section_no":     "1",
        "exam_date":      "2026-03-23",
        "exam_time":      "12.00-15.00",
        "room_name":      "PSB 1101",
        "teacher_name":   "อ.ดร.ศิรดา เขมานิฏฐาไท",
        "supervisors":    [
            {"name": "อ.ดร.ศิรดา เขมานิฏฐาไท", "slot_order": 1},
            {"name": "นางสาวอิสรีย์ พงศ์ทิวาพรรณ", "slot_order": 2},
        ],
        "num_students":   83,
        "num_exam_sets":  83,
        "exam_type":      "final",
        "semester":       "2",
        "academic_year":  "2568",
    }

    # นักศึกษาตัวอย่าง
    sample_students = [
        {"student_id": f"68191{str(i).zfill(4)}", "student_name": f"นักศึกษา ตัวอย่าง{i}"}
        for i in range(1, 84)
    ]

    out_dir = "/home/claude/doc_output"
    os.makedirs(out_dir, exist_ok=True)

    qr_path = "/home/claude/image1.png" if os.path.exists("/home/claude/image1.png") else None

    # 1. ใบปะหน้า
    data = generate_cover_page(**{k: SAMPLE[k] for k in
        ["course_id","course_name_th","section_no","exam_date","exam_time","room_name","exam_type","semester","academic_year"]})
    with open(f"{out_dir}/ใบปะหน้า_{SAMPLE['course_id']}_{SAMPLE['section_no']}_{SAMPLE['room_name']}.docx", "wb") as f:
        f.write(data)
    print("✅ ใบปะหน้า")

    # 2. ปกซอง
    data = generate_envelope_cover(
        **{k: SAMPLE[k] for k in ["course_id","section_no","exam_date","exam_time","room_name",
                                   "teacher_name","supervisors","num_students","num_exam_sets",
                                   "exam_type","semester","academic_year"]},
        qr_image_path=qr_path
    )
    with open(f"{out_dir}/ปกซอง_{SAMPLE['course_id']}_{SAMPLE['section_no']}.docx", "wb") as f:
        f.write(data)
    print("✅ ปกซอง")

    # 3. ใบลงมือชื่อ
    data = generate_attendance_sheet(
        **{k: SAMPLE[k] for k in ["course_id","section_no","exam_date","exam_time","room_name",
                                   "teacher_name","exam_type","semester","academic_year"]},
        students=sample_students
    )
    with open(f"{out_dir}/ใบลงมือชื่อ_{SAMPLE['course_id']}_{SAMPLE['section_no']}.docx", "wb") as f:
        f.write(data)
    print("✅ ใบลงมือชื่อ")

    print(f"\nOutput: {out_dir}/")
